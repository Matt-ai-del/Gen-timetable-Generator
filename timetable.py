import streamlit as st
import pandas as pd
import random
from collections import defaultdict
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime, timedelta
import multiprocessing
from concurrent.futures import ProcessPoolExecutor
import json
import logging
import uuid

from constants import DAYS, TIME_SLOTS, WEEKLY_MODULE_HOURS, SESSIONS_PER_MODULE

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Constants with 4 time slots
MODULE_MAX_DAILY_HOURS = 4  # Max 4 hours/day (2 sessions)
MODULE_MAX_WEEKLY_HOURS = 8  # Max 8 hours/week (4 sessions)
FACULTY_MIN_HOURS = 8
FACULTY_MAX_HOURS = 40  # 20 sessions (4 slots/day × 5 days) × 2 hours = 40 hours max
FACULTY_DAILY_MAX_HOURS = 8  # Max 8 hours/day (4 sessions)
POPULATION_SIZE = 200  # Optimized for faster execution
GENERATIONS = 1000     # Optimized for faster convergence
MUTATION_RATE = 0.15   # Increased for better exploration
NUM_PARALLEL_PROCESSES = max(1, multiprocessing.cpu_count() - 1)

def initialize_population(data):
    "Initialize a population of timetables with better distribution"
    population = []
    max_attempts = 200  # for larger module count
    
    for i in range(POPULATION_SIZE):
        try:
            timetable = {
                'slots': {day: {slot: [] for slot in TIME_SLOTS} for day in DAYS},
                'lecturer_slots': {l['id']: {day: {slot: None for slot in TIME_SLOTS} for day in DAYS} for l in data['lecturers']},
                'room_slots': {r['name'] if isinstance(r, dict) else r: {day: {slot: None for slot in TIME_SLOTS} for day in DAYS} for r in data['rooms']}
            }
            
            # Track assigned modules and their daily session counts
            assigned_modules = set()
            module_daily_sessions = defaultdict(lambda: defaultdict(int))
            
            # Group modules by target groups for better distribution
            modules_by_group = defaultdict(list)
            for module in data['modules']:
                if not module['code']:
                    continue
                for (program, level) in module.get('target_groups', []):
                    modules_by_group[(program, level)].append(module)
            
            # Sort modules by hours and group size (descending)
            for group_key, modules in modules_by_group.items():
                # Sort by hours first (fixed requirement) then by number of target groups
                modules.sort(key=lambda x: (-x.get('hours', WEEKLY_MODULE_HOURS), -len(x.get('target_groups', []))))
                
                for module in modules:
                    if module['code'] in assigned_modules:
                        continue
                    
                    # Use the module's hours to determine required sessions (2 hours per session)
                    required_sessions = module.get('hours', WEEKLY_MODULE_HOURS) // 2
                    
                    # Track attempts to schedule this module
                    scheduling_attempts = 0
                    while scheduling_attempts < max_attempts and module['code'] not in assigned_modules:
                        scheduling_attempts += 1
                        
                        # Find available slots
                        available_slots = []
                        for day in DAYS:
                            # Skip if module already has a session on this day
                            if module_daily_sessions[module['code']][day] >= 1:
                                continue
                            
                            for slot in TIME_SLOTS:
                                # Increase slot capacity to 10 modules
                                if len(timetable['slots'][day][slot]) >= 10:
                                    continue
                                
                                # Check for student group conflicts - only check actual overlaps
                                has_conflict = False
                                module_groups = set(module.get('target_groups', []))
                                for entry in timetable['slots'][day][slot]:
                                    other_module = next((m for m in data['modules'] if m['code'] == entry['module']), None)
                                    if other_module:
                                        other_groups = set(other_module.get('target_groups', []))
                                        if module_groups.intersection(other_groups):
                                            has_conflict = True
                                            break
                                if has_conflict:
                                    continue
                                
                                # Find available lecturer
                                available_lecturers = []
                                for lecturer in data['lecturers']:
                                    if module['code'] not in lecturer.get('modules', []):
                                        continue
                                    
                                    # Check if lecturer is already teaching in this slot
                                    lecturer_busy = False
                                    for entry in timetable['slots'][day][slot]:
                                        if entry['lecturer'] == lecturer['name']:
                                            lecturer_busy = True
                                            break
                                    
                                    if lecturer_busy:
                                        continue
                                    
                                    # Check lecturer's daily hours
                                    lecturer_daily_hours = sum(1 for s in TIME_SLOTS 
                                                             if any(e['lecturer'] == lecturer['name'] 
                                                                    for e in timetable['slots'][day][s]))
                                    if lecturer_daily_hours >= 8:  # Maximum 8 hours per day
                                        continue
                                    
                                    available_lecturers.append(lecturer)
                                
                                if not available_lecturers:
                                    continue
                                
                                # Find available room with sufficient capacity
                                available_rooms = []
                                for room in data['rooms']:
                                    room_name = room['name'] if isinstance(room, dict) else room
                                    room_capacity = room['capacity'] if isinstance(room, dict) else 500
                                    
                                    # Check if room is already booked
                                    if any(e['room'] == room_name for e in timetable['slots'][day][slot]):
                                        continue
                                    
                                    # Calculate total students for this slot
                                    total_students = sum(
                                        data.get('students_per_group', {}).get((program, level), 0)
                                        for (program, level) in module.get('target_groups', [])
                                    )
                                    
                                    # Add students from other modules in this slot
                                    for entry in timetable['slots'][day][slot]:
                                        other_module = next((m for m in data['modules'] if m['code'] == entry['module']), None)
                                        if other_module:
                                            for (program, level) in other_module.get('target_groups', []):
                                                total_students += data.get('students_per_group', {}).get((program, level), 0)
                                    
                                    if total_students > room_capacity:
                                        continue
                                    
                                    available_rooms.append(room)
                                
                                if available_rooms:
                                    available_slots.append((day, slot, available_lecturers, available_rooms))
                        
                        if not available_slots:
                            logger.warning(f"No available slots found for module {module['code']} on attempt {scheduling_attempts}")
                            continue
                        
                        # Select random slot
                        day, slot, lecturers, rooms = random.choice(available_slots)
                        lecturer = random.choice(lecturers)
                        room = random.choice(rooms)
                        
                        # Add to timetable
                        timetable['slots'][day][slot].append({
                            'module': module['code'],
                            'lecturer': lecturer['name'],
                            'room': room['name'] if isinstance(room, dict) else room
                        })
                        module_daily_sessions[module['code']][day] += 1
                        
                        # Update lecturer and room slots
                        timetable['lecturer_slots'][lecturer['id']][day][slot] = module['code']
                        timetable['room_slots'][room['name'] if isinstance(room, dict) else room][day][slot] = module['code']
                        
                        # Check if we've scheduled all required sessions
                        total_sessions = sum(module_daily_sessions[module['code']].values())
                        if total_sessions >= required_sessions:
                            assigned_modules.add(module['code'])
                            break
            
            # Validate timetable before adding to population
            if timetable and 'slots' in timetable:
                # Verify all hard constraints are satisfied
                if validate_hard_constraints(timetable, data):
                    population.append(timetable)
                else:
                    logger.error(f"Timetable {i} failed hard constraint validation")
            else:
                logger.error(f"Invalid timetable generated in iteration {i}")
                continue
                
        except Exception as e:
            logger.error(f"Error initializing timetable {i}: {str(e)}")
            continue
    
    if not population:
        raise ValueError("Failed to initialize any valid timetables")
    
    return population

def validate_hard_constraints(timetable, data):
    "Validate that all hard constraints are strictly satisfied"
    try:
        # 1. Check student group conflicts - NO EXCEPTIONS
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if not isinstance(slot_entries, list):
                    continue
                    
                # Track which groups are scheduled in this slot
                scheduled_groups = set()
                for entry in slot_entries:
                    module = next((m for m in data['modules'] if m['code'] == entry['module']), None)
                    if not module:
                        logger.error(f"Invalid module code: {entry['module']}")
                        return False
                        
                    for group in module.get('target_groups', []):
                        if group in scheduled_groups:
                            logger.error(f"CRITICAL: Student group conflict - {group} has multiple classes in {day} {slot}")
                            return False
                        scheduled_groups.add(group)
        
        # 2. Check lecturer conflicts - NO EXCEPTIONS
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if not isinstance(slot_entries, list):
                    continue
                    
                # Track which lecturers are teaching in this slot
                scheduled_lecturers = set()
                for entry in slot_entries:
                    if entry['lecturer'] in scheduled_lecturers:
                        logger.error(f"CRITICAL: Lecturer conflict - {entry['lecturer']} teaching multiple classes in {day} {slot}")
                        return False
                    scheduled_lecturers.add(entry['lecturer'])
                    
                    # Check lecturer's daily hours
                    lecturer_daily_hours = sum(1 for s in TIME_SLOTS 
                                             if any(e['lecturer'] == entry['lecturer'] 
                                                    for e in timetable['slots'][day][s]))
                    if lecturer_daily_hours > FACULTY_DAILY_MAX_HOURS:
                        logger.error(f"CRITICAL: Lecturer daily overload - {entry['lecturer']} exceeds {FACULTY_DAILY_MAX_HOURS} hours on {day}")
                        return False
        
        # 3. Check room conflicts - NO EXCEPTIONS
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if not isinstance(slot_entries, list):
                    continue
                    
                # Track which rooms are used in this slot
                scheduled_rooms = set()
                for entry in slot_entries:
                    if entry['room'] in scheduled_rooms:
                        logger.error(f"CRITICAL: Room conflict - {entry['room']} double-booked in {day} {slot}")
                        return False
                    scheduled_rooms.add(entry['room'])
        
        # 4. Check room capacity - NO EXCEPTIONS
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if not isinstance(slot_entries, list):
                    continue
                    
                for entry in slot_entries:
                    room = next((r for r in data['rooms'] if (isinstance(r, dict) and r['name'] == entry['room']) or r == entry['room']), None)
                    if not room:
                        logger.error(f"CRITICAL: Invalid room - {entry['room']} not found in room list")
                        return False
                        
                    capacity = room['capacity'] if isinstance(room, dict) else 30
                    module = next((m for m in data['modules'] if m['code'] == entry['module']), None)
                    if not module:
                        logger.error(f"CRITICAL: Invalid module - {entry['module']} not found in module list")
                        return False
                        
                    total_students = sum(
                        data.get('students_per_group', {}).get((program, level), 0)
                        for (program, level) in module.get('target_groups', [])
                    )
                    
                    if total_students > capacity:
                        logger.error(f"CRITICAL: Room capacity exceeded - {entry['room']} has {total_students} students (capacity: {capacity})")
                        return False
        
        # 5. Check lecturer qualifications - NO EXCEPTIONS
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if not isinstance(slot_entries, list):
                    continue
                    
                for entry in slot_entries:
                    lecturer = next((l for l in data['lecturers'] if l['name'] == entry['lecturer']), None)
                    if not lecturer:
                        logger.error(f"CRITICAL: Invalid lecturer - {entry['lecturer']} not found in lecturer list")
                        return False
                        
                    if entry['module'] not in lecturer.get('modules', []):
                        logger.error(f"CRITICAL: Lecturer not qualified - {entry['lecturer']} not qualified to teach {entry['module']}")
                        return False
        
        # 6. Check module hours - NO EXCEPTIONS
        for module in data['modules']:
            if not module['code']:
                continue
                
            # Count scheduled sessions
            scheduled_sessions = 0
            daily_counts = defaultdict(int)
            for day in DAYS:
                for slot in TIME_SLOTS:
                    slot_entries = timetable['slots'][day][slot]
                    if isinstance(slot_entries, list):
                        if any(t and t['module'] == module['code'] for t in slot_entries):
                            scheduled_sessions += 1
                            daily_counts[day] += 1
                    elif isinstance(slot_entries, dict) and slot_entries:
                        if slot_entries['module'] == module['code']:
                            scheduled_sessions += 1
                            daily_counts[day] += 1
            
            required_sessions = module.get('hours', WEEKLY_MODULE_HOURS) // 2
            if scheduled_sessions != required_sessions:
                logger.error(
                    f"CRITICAL: Module hours mismatch - {module['code']} has {scheduled_sessions*2}h scheduled "
                    f"(required: {WEEKLY_MODULE_HOURS}h)"
                )
                return False

            for day, count in daily_counts.items():
                if count > 1:
                    logger.error(f"CRITICAL: Module {module['code']} is scheduled more than once on {day}")
                    return False
        
        return True
        
    except Exception as e:
        logger.error(f"CRITICAL: Error validating hard constraints: {str(e)}")
        return False

def calculate_fitness(timetable, data):
    "Calculate fitness score for a timetable"
    score = 1000  # Start with a base score
    
    # Track module distribution
    module_distribution = defaultdict(list)
    module_daily_counts = defaultdict(lambda: defaultdict(int))
    slot_usage = defaultdict(int)  # Track usage of each time slot
    
    # Check module distribution
    for module in data['modules']:
        if not module['code']:  # Skip empty modules
            continue
           
        module_slots = []
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if isinstance(slot_entries, list):
                    if any(t and t['module'] == module['code'] for t in slot_entries):
                        module_slots.append((day, slot))
                        module_distribution[module['code']].append((day, slot))
                        module_daily_counts[module['code']][day] += 1
                        slot_usage[slot] += 1
                elif isinstance(slot_entries, dict) and slot_entries:
                    if slot_entries['module'] == module['code']:
                        module_slots.append((day, slot))
                        module_distribution[module['code']].append((day, slot))
                        module_daily_counts[module['code']][day] += 1
                        slot_usage[slot] += 1
        
        # Each module must run for exactly two sessions (4 hours)
        required_sessions = SESSIONS_PER_MODULE
            
        # Penalize if module sessions don't match required sessions
        if len(module_slots) != required_sessions:
            score -= abs(len(module_slots) - required_sessions) * 100
        
        # Penalize if module is scheduled on same day
        days_with_module = set(day for day, _ in module_slots)
        if len(days_with_module) < len(module_slots):
            score -= 20  # Penalty for same-day scheduling
        
        # Penalize if module sessions are too close together
        if len(module_slots) > 1:
            for i in range(len(module_slots)-1):
                day1, slot1 = module_slots[i]
                day2, slot2 = module_slots[i+1]
                if day1 == day2:
                    slot1_idx = TIME_SLOTS.index(slot1)
                    slot2_idx = TIME_SLOTS.index(slot2)
                    if abs(slot1_idx - slot2_idx) == 1:  # Consecutive slots
                        score -= 15  # Penalty for consecutive slots
    
    # Penalize uneven distribution of modules across days
    day_module_counts = defaultdict(int)
    for module_slots in module_distribution.values():
        for day, _ in module_slots:
            day_module_counts[day] += 1
    
    if day_module_counts:
        max_count = max(day_module_counts.values())
        min_count = min(day_module_counts.values())
        score -= (max_count - min_count) * 10  # Penalty for uneven distribution
    
    # Penalize uneven distribution across time slots
    if slot_usage:
        max_slot_usage = max(slot_usage.values())
        min_slot_usage = min(slot_usage.values())
        score -= (max_slot_usage - min_slot_usage) * 5  # Penalty for uneven slot usage
    
    # Check room capacity violations
    for day in DAYS:
        for slot in TIME_SLOTS:
            slot_entries = timetable['slots'][day][slot]
            if isinstance(slot_entries, list):
                entries = slot_entries
            elif isinstance(slot_entries, dict) and slot_entries:
                entries = [slot_entries]
            else:
                entries = []
            
            room_usage = defaultdict(int)
            for entry in entries:
                room = entry['room']
                room_usage[room] += 1
                if room_usage[room] > 1:
                    score -= 50  # Penalty for room double-booking
    
    return score

def crossover(parent1, parent2, data):
    "Perform crossover between two timetables"
    child = {
        'slots': {day: {slot: [] for slot in TIME_SLOTS} for day in DAYS},
        'lecturer_slots': {l["id"]: {day: {slot: None for slot in TIME_SLOTS} for day in DAYS} 
            for l in data['lecturers']},
        'room_slots': {room['name'] if isinstance(room, dict) else room: {day: {slot: None for slot in TIME_SLOTS} for day in DAYS} 
            for room in data['rooms']}
    }
    
    # Pre-calculate module requirements based on module hours
    module_requirements = {
        module['code']: module.get('hours', WEEKLY_MODULE_HOURS) // 2
        for module in data['modules']
        if module['code']
    }
    
    # Track assigned sessions for each module
    assigned_sessions = defaultdict(int)
    
    # Randomly select days from each parent
    for day in DAYS:
        if random.random() < 0.5:
            source = parent1
        else:
            source = parent2
            
        # Copy slots from source parent
        for slot in TIME_SLOTS:
            slot_entries = source['slots'][day][slot]
            if isinstance(slot_entries, list):
                entries = slot_entries
            elif isinstance(slot_entries, dict) and slot_entries:
                entries = [slot_entries]
            else:
                entries = []
            
            for entry in entries:
                module_code = entry['module']
                # Check if we can add more sessions for this module
                if assigned_sessions[module_code] < module_requirements[module_code]:
                    # Verify no conflicts
                    if not any(t['module'] == module_code for t in child['slots'][day][slot]):
                        # Check lecturer availability
                        lecturer = next((l for l in data['lecturers'] if l['name'] == entry['lecturer']), None)
                        if lecturer and not child['lecturer_slots'][lecturer['id']][day][slot]:
                            # Check room availability
                            room = entry['room']
                            if not child['room_slots'][room][day][slot]:
                                # Add the entry
                                child['slots'][day][slot].append(entry)
                                child['lecturer_slots'][lecturer['id']][day][slot] = module_code
                                child['room_slots'][room][day][slot] = module_code
                                assigned_sessions[module_code] += 1
    
    return child

def mutate(timetable, data):
    "Apply random mutations to a timetable"
    # Select a random day and slot
    day = random.choice(DAYS)
    slot = random.choice(TIME_SLOTS)
    
    # Clear the selected slot
    timetable['slots'][day][slot] = []
    for l in data['lecturers']:
        timetable['lecturer_slots'][l['id']][day][slot] = None
    for room in data['rooms']:
        room_name = room['name'] if isinstance(room, dict) else room
        timetable['room_slots'][room_name][day][slot] = None
    
    # Find a module that needs rescheduling
    available_modules = []
    for module in data['modules']:
        if not module['code']:
            continue
            
        # Count current sessions
        current_sessions = 0
        for d in DAYS:
            for s in TIME_SLOTS:
                slot_entries = timetable['slots'][d][s]
                if isinstance(slot_entries, list):
                    if any(t and t['module'] == module['code'] for t in slot_entries):
                        current_sessions += 1
                elif isinstance(slot_entries, dict) and slot_entries:
                    if slot_entries['module'] == module['code']:
                        current_sessions += 1
        
        required_sessions = module.get('hours', WEEKLY_MODULE_HOURS) // 2
        
        if current_sessions < required_sessions:
            available_modules.append(module)
    
    if not available_modules:
        return
    
    # Select a random module to reschedule
    module = random.choice(available_modules)
    
    # Check daily limit
    per_day_count = sum(1 for s in TIME_SLOTS 
                       if any(t and t['module'] == module['code'] 
                             for t in timetable['slots'][day][s]))
    if per_day_count >= 1:  # Max 1 session (2 hours) per day for any module
        return
    
    # Check for conflicts
    if any(t['module'] == module['code'] for t in timetable['slots'][day][slot]):
        return
    
    # Check for student group conflicts
    group_clash = False
    for (program, level) in module.get('target_groups', []):
        for t in timetable['slots'][day][slot]:
            tmod = next((m for m in data['modules'] if m['code'] == t['module']), None)
            if tmod and (program, level) in tmod.get('target_groups', []):
                group_clash = True
                break
    if group_clash:
        return
    
    # Find available lecturer
    available_lecturers = [
        l for l in data['lecturers']
        if module['code'] in l['modules']
        and not timetable['lecturer_slots'][l['id']][day][slot]
        and sum(1 for s in TIME_SLOTS if timetable['lecturer_slots'][l['id']][day][s]) < 4
    ]
    if not available_lecturers:
        return
    
    lecturer = random.choice(available_lecturers)
    
    # Find available room
    available_rooms = []
    
    # Get all programs for this module
    module_programs = {program for program, _ in module.get('target_groups', [])}
    
    for room in data['rooms']:
        # Handle both dictionary and string room formats
        is_dict = isinstance(room, dict)
        room_name = room['name'] if is_dict else room
        
        # Skip if room is already booked for this slot
        if timetable['room_slots'][room_name][day][slot]:
            continue
        
        # If room has program restrictions, enforce them strictly
        if is_dict and room.get('allowed_programs'):
            # If module has no programs, it can't be in a restricted room
            if not module_programs:
                continue
                
            # Check if ALL of the module's programs are allowed in this room
            allowed = all(program in room['allowed_programs'] for program in module_programs)
            if not allowed:
                continue
        # If room has no program restrictions, it can only be used by modules with no programs
        elif module_programs:
            continue
        
        # Calculate total students for this module's groups
        total_students = 0
        for program, level in module.get('target_groups', []):
            total_students += data.get('students_per_group', {}).get((program, level), 0)
        
        # Add students from other modules in this slot
        for t in timetable['slots'][day][slot]:
            tmod = next((m for m in data['modules'] if m['code'] == t['module']), None)
            if tmod:
                for program, level in tmod.get('target_groups', []):
                    total_students += data.get('students_per_group', {}).get((program, level), 0)
        
        # Get room capacity, default to 30 for old format
        capacity = room['capacity'] if is_dict else 500
        
        # Check if room can accommodate all students
        if total_students <= capacity:
            available_rooms.append(room_name)
    
    if not available_rooms:
        return
    
    room_name = random.choice(available_rooms)
    
    # Add the new entry
    slot_entry = {
        'module': module['code'],
        'lecturer': lecturer['name'],
        'room': room_name
    }
    timetable['slots'][day][slot].append(slot_entry)
    timetable['lecturer_slots'][lecturer['id']][day][slot] = module['code']
    timetable['room_slots'][room_name][day][slot] = module['code']

def tournament_selection(population, fitness_scores, tournament_size=3):
    "Select a parent using tournament selection"
    # Ensure we have valid indices
    valid_indices = [i for i in range(len(population)) if i < len(fitness_scores) and population[i] is not None]
    
    if not valid_indices:
        logger.error("No valid individuals available for tournament selection")
        # Return a random valid timetable from the population
        valid_timetables = [t for t in population if t is not None]
        if not valid_timetables:
            raise ValueError("No valid timetables available for selection")
        return random.choice(valid_timetables)
    
    # Randomly select tournament_size individuals from valid indices
    tournament_indices = random.sample(valid_indices, min(tournament_size, len(valid_indices)))
    
    # Get their fitness scores
    tournament_fitness = [fitness_scores[i] for i in tournament_indices]
    
    # Select the best one
    winner_idx = tournament_indices[tournament_fitness.index(max(tournament_fitness))]
    
    return population[winner_idx]

def parallel_fitness_calculation(timetable_data):
    "Calculate fitness for a timetable in parallel"
    timetable, data = timetable_data
    return calculate_fitness(timetable, data)

def save_lecturer_timetable(lecturer_name, timetable, department, session_title):
    "Save a lecturer's timetable to the database"
    try:
        import sqlite3
        from datetime import datetime
        
        # Connect to the database
        conn = sqlite3.connect('timetables.db')
        cursor = conn.cursor()
        
        # Create lecturer_timetables table if it doesn't exist
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS lecturer_timetables (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                lecturer_name TEXT NOT NULL,
                department TEXT NOT NULL,
                session_title TEXT NOT NULL,
                timetable_data TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Convert timetable to JSON string
        import json
        timetable_json = json.dumps(timetable)
        
        # Insert the timetable
        cursor.execute('''
            INSERT INTO lecturer_timetables (lecturer_name, department, session_title, timetable_data)
            VALUES (?, ?, ?, ?)
        ''', (lecturer_name, department, session_title, timetable_json))
        
        # Commit the changes
        conn.commit()
        logger.info(f"Saved timetable for lecturer {lecturer_name}")
        
    except Exception as e:
        logger.error(f"Error saving timetable for lecturer {lecturer_name}: {str(e)}")
        raise
    finally:
        if 'conn' in locals():
            conn.close()

def generate_timetable(data):
    "Generate a timetable using genetic algorithm with progress tracking"
    logger.info("Starting timetable generation...")
    
    # Initialize progress tracking
    progress_bar = st.progress(0)
    progress_text = st.empty()
    fitness_chart = st.empty()
    fitness_scores = []
    
    try:
        # Initialize population
        logger.info("Initializing population...")
        population = initialize_population(data)
        if not population:
            logger.error("Failed to initialize population")
            return None
        
        best_fitness = float('-inf')
        best_timetable = None
        generations_without_improvement = 0
        
        # Main genetic algorithm loop
        for generation in range(GENERATIONS):
            # Update progress
            progress = (generation + 1) / GENERATIONS
            progress_bar.progress(progress)
            progress_text.text(f"Generation {generation + 1}/{GENERATIONS}")
            
            # Calculate fitness scores
            fitness_scores_batch = []
            for timetable in population:
                try:
                    fitness = calculate_fitness(timetable, data)
                    fitness_scores_batch.append(fitness)
                except Exception as e:
                    logger.error(f"Error calculating fitness: {e}")
                    fitness_scores_batch.append(float('-inf'))
            
            # Update best solution
            max_fitness = max(fitness_scores_batch)
            fitness_scores.append(max_fitness)
            
            # Update fitness chart
            fitness_chart.line_chart(fitness_scores)
            
            if max_fitness > best_fitness:
                best_fitness = max_fitness
                best_timetable = population[fitness_scores_batch.index(max_fitness)]
                generations_without_improvement = 0
            else:
                generations_without_improvement += 1
            
            # Early stopping check
            if generations_without_improvement >= 20:
                logger.info(f"Early stopping at generation {generation + 1}")
                break
            
            # Selection and reproduction
            new_population = []
            while len(new_population) < POPULATION_SIZE:
                parent1 = tournament_selection(population, fitness_scores_batch)
                parent2 = tournament_selection(population, fitness_scores_batch)
                child = crossover(parent1, parent2, data)
                if random.random() < MUTATION_RATE:
                    mutate(child, data)
                new_population.append(child)
        
        population = new_population
    
        # Final progress update
        progress_bar.progress(1.0)
        progress_text.text("Timetable generation complete!")
        
        if best_timetable:
            logger.info("Timetable generation successful")
            return best_timetable
        else:
            logger.error("Failed to generate valid timetable")
            return None
    
    except Exception as e:
        logger.error(f"Error in timetable generation: {e}")
        return None
    
    finally:
        # Clean up progress indicators
        progress_bar.empty()
        progress_text.empty()
        fitness_chart.empty()

    # Check for modules not assigned to any lecturer
    unassigned_modules = [m['code'] for m in data['modules'] if not any(l for l in data['lecturers'] if m['code'] in l.get('modules', []))]
    if unassigned_modules:
        st.error("Timetable generation failed because the following modules are not assigned to any lecturer:")
        for code in unassigned_modules:
            st.write(f"- {code}")
    else:
        violations = validate_timetable(best_timetable, data)
        if violations:
            st.error("Timetable generation failed for the following reasons:")
            for v in violations:
                st.write(f"- {v}")
        else:
            st.success("Timetable generated successfully!")

def display_timetable(timetable, data, department, programs, levels, session_title):
    # Build dynamic header
    programs_str = '/'.join(programs)
    levels_str = ', '.join(levels)
    header_html = f"""
    <h3 style='text-align: center;'>DEPARTMENT OF {department.upper()}</h3>
    <h4 style='text-align: center;'>{programs_str} LEVEL {levels_str} <br>{session_title}</h4>
    """
    st.markdown(header_html, unsafe_allow_html=True)

    # Build table: days as rows, time slots as columns
    table = []
    for day in DAYS:
        row = []
        for slot in TIME_SLOTS:
            slot_entries = timetable['slots'][day][slot]
            # Group by (module_code, room)
            grouped = {}
            for entry in slot_entries:
                module_code = entry['module']
                room = entry['room']
                key = (module_code, room)
                if key not in grouped:
                    grouped[key] = []
                module = next((m for m in data['modules'] if m['code'] == module_code), None)
                if module:
                    for (program, level) in module.get('target_groups', []):
                        grouped[key].append(f"{level} {program}")
            cell_lines = []
            for (module_code, room), group_list in grouped.items():
                group_str = '/'.join(sorted(set(group_list)))
                cell_lines.append(f"{module_code} {group_str} ({room})")
            # Empty cell with consistent size
            cell_content = '\n'.join(cell_lines) if cell_lines else ""
            row.append(cell_content)
        table.append(row)

    df = pd.DataFrame(table, index=DAYS, columns=TIME_SLOTS)
    # Set a fixed width and height for all cells
    st.dataframe(
        df.style.set_properties(**{
            'white-space': 'pre',
            'min-width': '150px',
            'max-width': '150px',
            'min-height': '100px',
            'max-height': '100px'
        }),
        use_container_width=True
    )

    # DOCX export
    doc = Document()
    doc.add_heading(f'DEPARTMENT OF {department.upper()}', level=1)
    doc.add_heading(f'{programs_str} LEVEL {levels_str}', level=2)
    doc.add_heading(session_title, level=2)

    # Add combined timetable: days as rows, time slots as columns
    table_doc = doc.add_table(rows=len(DAYS)+1, cols=len(TIME_SLOTS)+1)
    table_doc.style = 'Table Grid'

    # Add headers
    header_cells = table_doc.rows[0].cells
    header_cells[0].text = "Day"
    for col_idx, slot in enumerate(TIME_SLOTS, 1):
        header_cells[col_idx].text = slot

    # Add data
    for row_idx, day in enumerate(DAYS, 1):
        row_cells = table_doc.rows[row_idx].cells
        row_cells[0].text = day
        for col_idx, slot in enumerate(TIME_SLOTS, 1):
            slot_entries = timetable['slots'][day][slot]
            # Group by (module_code, room)
            grouped = {}
            for entry in slot_entries:
                module_code = entry['module']
                room = entry['room']
                key = (module_code, room)
                if key not in grouped:
                    grouped[key] = []
                module = next((m for m in data['modules'] if m['code'] == module_code), None)
                if module:
                    for (program, level) in module.get('target_groups', []):
                        grouped[key].append(f"{level} {program}")
            cell_lines = []
            for (module_code, room), group_list in grouped.items():
                group_str = '/'.join(sorted(set(group_list)))
                cell_lines.append(f"{module_code} {group_str} ({room})")
            # Empty cell for consistency
            cell_content = '\n'.join(cell_lines) if cell_lines else ""
            row_cells[col_idx].text = cell_content

    # Save to bytes buffer
    # Add grouped appendix by program and level
    doc.add_page_break()
    doc.add_heading(f'Module and Lecturer Assignments for {department} Department', level=2)
    
    # Program name mapping
    program_names = {
        'CS': 'Computer Science',
        'CSE': 'Computer Systems Engineering',
        'SWE': 'Software Engineering',
        'CSEC': 'Computer Security'
    }
    
    # Get all unique (program, level) pairs from modules
    program_level_pairs = set()
    for module in data['modules']:
        for (program, level) in module.get('target_groups', []):
            program_level_pairs.add((program, level))
    
    # Sort for consistent order
    for program, level in sorted(program_level_pairs, key=lambda x: (x[0], x[1])):
        full_program_name = program_names.get(program, program)  # Use mapping or fallback to code
        doc.add_paragraph(f"{full_program_name} Level {level}", style='List Bullet')
        for module in data['modules']:
            if (program, level) in module.get('target_groups', []):
                # Find all lecturers who teach this module
                lecturers = [lect['name'] for lect in data['lecturers'] if module['code'] in lect.get('modules', [])]
                lecturers_str = ', '.join(lecturers) if lecturers else ''
                doc.add_paragraph(f"{module['code']}  {module['name']}{' ' * (40 - len(module['code'] + module['name']))}{lecturers_str}", style='List Continue')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Download button
    st.download_button(
        label="Download Combined Timetable (DOCX)",
        data=buffer,
        file_name=f"{department}_combined_timetable.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_combined_timetable_{uuid.uuid4()}",
        use_container_width=True
    )
    
    # Module Statistics
    st.header("Module Statistics")
    course_data = []
    scheduled_hours = defaultdict(int)
    
    for day in DAYS:
        for slot in TIME_SLOTS:
            slot_data = timetable['slots'][day][slot]
            if isinstance(slot_data, list):
                entries = slot_data
            elif isinstance(slot_data, dict) and slot_data:
                entries = [slot_data]
            else:
                entries = []
            
            for entry in entries:
                if entry:
                    course = entry['module']
                    scheduled_hours[course] += 2  # Each slot is 2 hours
    
    for course in data['modules']:
        course_data.append({
            "Course Code": str(course["code"]),
            "Course Name": str(course["name"]),
            "Required Hours": int(course["hours"]),
            "Scheduled Hours": int(scheduled_hours[course["code"]]),
            "Status": "✅" if scheduled_hours[course["code"]] == course["hours"] else "❌"
        })
    
    # Create DataFrame with explicit dtypes
    course_df = pd.DataFrame(course_data)
    # Ensure all columns have consistent types
    course_df = course_df.astype({
        'Course Code': 'string',
        'Course Name': 'string',
        'Required Hours': 'int32',
        'Scheduled Hours': 'int32',
        'Status': 'string'
    })
    st.dataframe(course_df, use_container_width=True)
    
    # Lecturers Statistics
    st.header("Lecturers Statistics")
    faculty_data = []
    faculty_hours = defaultdict(int)
    for day in DAYS:
        for slot in TIME_SLOTS:
            slot_data = timetable['slots'][day][slot]
            if isinstance(slot_data, list):
                entries = slot_data
            elif isinstance(slot_data, dict) and slot_data:
                entries = [slot_data]
            else:
                entries = []
            for entry in entries:
                if entry:
                    for faculty in data['lecturers']:
                        if faculty['name'] == entry['lecturer']:
                            faculty_hours[faculty['id']] += 2  # Each slot is 2 hours
    # Updated status logic
    faculty_data = []
    for faculty in data['lecturers']:
        hours = faculty_hours[faculty['id']]
        # Required hours is the sum of hours for all modules this lecturer is assigned to
        required_hours = sum(
            m['hours'] for m in data['modules'] if m['code'] in faculty.get('modules', [])
        )
        if hours == required_hours:
            status = "✅"
        elif hours < required_hours:
            status = "❌"
        else:
            status = "⚠️"
        faculty_data.append({
            "Name": str(faculty["name"]),
            "Total Hours": int(hours),
            "Required Hours": int(required_hours),
            "Status": str(status)
        })
    
    # Create DataFrame with explicit dtypes
    faculty_df = pd.DataFrame(faculty_data)
    # Ensure all columns have consistent types
    faculty_df = faculty_df.astype({
        'Name': 'string',
        'Total Hours': 'int32',
        'Required Hours': 'int32',
        'Status': 'string'
    })
    st.dataframe(faculty_df, use_container_width=True)

    st.markdown(
        """
        <hr style='margin-top:2em;margin-bottom:0.5em;border:1px solid #eee;'>
        <div style='text-align:center; color: #888; font-size: 1.1em;'>
            Empowering Education Through Technology | Midlands State University
        </div>
        """,
        unsafe_allow_html=True
    )

def synchronize_shared_modules(university_data):
    "Detect modules that appear in more than one department and assign them the same (day, slot) across all departments. Modifies the university_data in-place to pre-fill shared module slots."
    from collections import defaultdict
    # 1. Build a mapping from module code to departments
    module_dept_map = defaultdict(list)
    for dept_name, dept_data in university_data['departments'].items():
        for module in dept_data['modules']:
            if module.get('code'):
                module_dept_map[module['code']].append(dept_name)

    # 2. Find shared modules
    shared_modules = {code: depts for code, depts in module_dept_map.items() if len(depts) > 1}
    if not shared_modules:
        return  # Nothing to do

    # 3. For each shared module, pick a (day, slot) and assign in all departments
    used_slots = set()
    for module_code, depts in shared_modules.items():
        assigned = False
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_taken = False
                for dept in depts:
                    dept_data = university_data['departments'][dept]
                    # Check if slot is already used for this module in this department
                    if 'preassigned_slots' not in dept_data:
                        dept_data['preassigned_slots'] = {}
                    for m in dept_data['modules']:
                        if m['code'] == module_code:
                            if (day, slot) in dept_data['preassigned_slots'].values():
                                slot_taken = True
                                break
                    # Check if slot is already used for any module in this department
                    for m in dept_data['modules']:
                        if 'assigned_slots' in m and (day, slot) in m['assigned_slots']:
                            slot_taken = True
                            break
                    if slot_taken:
                        break
                if not slot_taken and (day, slot) not in used_slots:
                    # Assign this slot to the module in all departments
                    for dept in depts:
                        dept_data = university_data['departments'][dept]
                        for m in dept_data['modules']:
                            if m['code'] == module_code:
                                if 'assigned_slots' not in m:
                                    m['assigned_slots'] = []
                                m['assigned_slots'].append((day, slot))
                        dept_data['preassigned_slots'][module_code] = (day, slot)
                    used_slots.add((day, slot))
                    assigned = True
                    break
            if assigned:
                break 

def validate_timetable(timetable, data):
    "Enhanced validation with detailed violation reports"
    # Track all violations
    violations = []
    reported_modules = set()  # Track which modules we've already reported
    
    # 1. Module constraints
    for module in data['modules']:
        if not module['code']:
            continue
            
        # Count scheduled sessions
        weekly_sessions = 0
        daily_counts = {day: 0 for day in DAYS}
        
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if isinstance(slot_entries, list):
                    if any(t and t['module'] == module['code'] for t in slot_entries):
                        weekly_sessions += 1
                        daily_counts[day] += 1
                elif isinstance(slot_entries, dict) and slot_entries:
                    if slot_entries['module'] == module['code']:
                        weekly_sessions += 1
                        daily_counts[day] += 1
        
        # Check weekly limit (8 hours = 4 sessions)
        required_sessions = module['hours'] // 2
        if weekly_sessions != required_sessions and module['code'] not in reported_modules:
            violations.append(
                f"MODULE SCHEDULING: {module['code']} has {weekly_sessions*2}h/week (should be {module['hours']}h)"
            )
            reported_modules.add(module['code'])
            continue  # Skip other checks for this module since we've already reported it
        
        # Only check daily limit if we haven't reported this module yet
        if module['code'] not in reported_modules:
            for day, count in daily_counts.items():
                if count > 2:
                    violations.append(
                        f"DAILY OVERLOAD: {module['code']} has {count*2}h on {day} (max 4h allowed)"
                    )
                    reported_modules.add(module['code'])
                    break  # Only report one daily overload per module

    # 2. Faculty constraints
    for lecturer in data['lecturers']:
        weekly_hours = 0
        daily_hours = {day: 0 for day in DAYS}
        
        for day in DAYS:
            for slot in TIME_SLOTS:
                slot_entries = timetable['slots'][day][slot]
                if isinstance(slot_entries, list):
                    if any(t and t['lecturer'] == lecturer['name'] for t in slot_entries):
                        daily_hours[day] += 2
                        weekly_hours += 2
                elif isinstance(slot_entries, dict) and slot_entries:
                    if slot_entries['lecturer'] == lecturer['name']:
                        daily_hours[day] += 2
                        weekly_hours += 2
        
        # Check weekly limit (40 hours)
        if weekly_hours > lecturer.get('max_weekly', 20)*2:
            violations.append(
                f"FACULTY OVERLOAD: {lecturer['name']} teaches {weekly_hours}h/week (max {lecturer.get('max_weekly', 20)*2}h allowed)"
            )
        
        # Check daily limit (8 hours)
        for day, hours in daily_hours.items():
            if hours > lecturer.get('max_daily', 4)*2:
                violations.append(
                    f"DAILY OVERLOAD: {lecturer['name']} teaches {hours}h on {day} (max {lecturer.get('max_daily', 4)*2}h allowed)"
                )

    # 3. Room constraints
    for day in DAYS:
        for slot in TIME_SLOTS:
            room_usage = {}
            for entry in timetable['slots'][day][slot]:
                # Get room name whether it's a string or dict
                room_name = entry['room']
                room = next((r for r in data['rooms'] if 
                           (isinstance(r, dict) and r['name'] == room_name) or 
                           r == room_name), None)
                
                if not room:
                    violations.append(f"INVALID ROOM: Room '{room_name}' not found in room list")
                    continue
                    
                # Check room capacity
                room_capacity = room['capacity'] if isinstance(room, dict) else 30
                module = next((m for m in data['modules'] if m['code'] == entry['module']), None)
                if module:
                    # Get module's target programs
                    module_programs = {program for program, _ in module.get('target_groups', [])}
                    
                    # Check if room has program restrictions
                    if isinstance(room, dict) and room.get('allowed_programs'):
                        # If module has no programs, it's invalid for this room
                        if not module_programs:
                            violations.append(
                                f"ROOM RESTRICTION: {entry['module']} has no assigned programs "
                                f"but is scheduled in {room_name} which has program restrictions"
                            )
                        # Check if any of module's programs are allowed in this room
                        elif not any(program in room['allowed_programs'] for program in module_programs):
                            allowed_programs = ", ".join(room['allowed_programs']) or 'None'
                            module_progs = ", ".join(module_programs) or 'None'
                            violations.append(
                                f"ROOM RESTRICTION: {entry['module']} (programs: {module_progs}) "
                                f"is not allowed in {room_name} (allowed: {allowed_programs})"
                            )
                    
                    # Check room capacity
                    total_students = sum(
                        data.get('students_per_group', {}).get((program, level), 0)
                        for (program, level) in module.get('target_groups', [])
                    )
                    if total_students > room_capacity:
                        violations.append(
                            f"ROOM CAPACITY: {entry['module']} in {room_name} has {total_students} students "
                            f"(capacity: {room_capacity})"
                        )

    return violations 
