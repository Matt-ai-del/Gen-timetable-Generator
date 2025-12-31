import sqlite3
import json
import multiprocessing
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
from collections import defaultdict
import streamlit as st
import logging
import threading
from queue import Queue, Empty
from typing import Optional, Dict, Any, List, Union
from contextlib import contextmanager
from constants import DAYS, TIME_SLOTS
from werkzeug.security import generate_password_hash
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Connection pool implementation
class ConnectionPool:
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        with cls._lock:
            if cls._instance is None:
                cls._instance = super(ConnectionPool, cls).__new__(cls)
                cls._instance._initialize_pool()
            return cls._instance
    
    def _initialize_pool(self):
        self.pool_size = min(20, multiprocessing.cpu_count() * 2)  # Dynamic pool size
        self._pool = Queue(maxsize=self.pool_size)
        self._in_use = {}
        self._lock = threading.Lock()
        
        for _ in range(self.pool_size):
            conn = sqlite3.connect('timetable.db', check_same_thread=False)
            # Performance optimizations
            conn.execute('PRAGMA journal_mode=WAL')  # Better concurrency
            conn.execute('PRAGMA busy_timeout=30000')  # 30 second timeout
            conn.execute('PRAGMA synchronous=NORMAL')  # Balance between safety and speed
            conn.execute('PRAGMA cache_size=10000')   # Increase cache for better performance
            conn.execute('PRAGMA temp_store=MEMORY')  # Store temp tables in memory
            self._pool.put(conn)
    
    @contextmanager
    def get_connection(self) -> sqlite3.Connection:
        conn = None
        try:
            conn = self._pool.get(timeout=30)  # 30 second timeout
            thread_id = threading.get_ident()
            self._in_use[thread_id] = conn
            yield conn
            conn.commit()
        except sqlite3.Error as e:
            if conn:
                conn.rollback()
            logging.error(f"Database error: {e}")
            raise
        except Exception as e:
            logging.error(f"Unexpected error: {e}")
            raise
        finally:
            if conn:
                thread_id = threading.get_ident()
                self._in_use.pop(thread_id, None)
                self._pool.put(conn)

# Initialize connection pool
db_pool = ConnectionPool()

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
TIME_SLOTS = ["08:00-10:00", "10:00-12:00", "12:00-14:00", "14:00-16:00"]

def convert_defaultdict_to_dict(obj):
    """Convert defaultdict and nested defaultdict objects to regular dictionaries"""
    try:
        if isinstance(obj, defaultdict):
            obj = {k: convert_defaultdict_to_dict(v) for k, v in obj.items()}
        elif isinstance(obj, dict):
            obj = {k: convert_defaultdict_to_dict(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            obj = [convert_defaultdict_to_dict(item) for item in obj]
        elif isinstance(obj, tuple):
            obj = tuple(convert_defaultdict_to_dict(item) for item in obj)
        elif hasattr(obj, '__dict__'):
            obj = convert_defaultdict_to_dict(obj.__dict__)
        return obj
    except Exception as e:
        logger.error(f"Error in convert_defaultdict_to_dict: {str(e)}")
        return str(obj)

def prepare_data_for_storage(data):
    """Prepare data for storage by converting all non-serializable objects"""
    try:
        # First convert defaultdicts to regular dicts
        data = convert_defaultdict_to_dict(data)
        
        # Remove any non-serializable objects
        if isinstance(data, dict):
            return {k: prepare_data_for_storage(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [prepare_data_for_storage(item) for item in data]
        elif isinstance(data, tuple):
            return tuple(prepare_data_for_storage(item) for item in data)
        elif isinstance(data, (str, int, float, bool, type(None))):
            return data
        elif isinstance(data, datetime):
            return data.isoformat()
        else:
            # For any other type, convert to string representation
            return str(data)
    except Exception as e:
        logger.error(f"Error preparing data for storage: {str(e)}")
        return str(data)

def safe_json_dumps(data):
    """Safely convert data to JSON string with error handling"""
    try:
        return json.dumps(data)
    except TypeError as e:
        logger.error(f"JSON serialization error: {str(e)}")
        # Try to identify the problematic data
        if isinstance(data, dict):
            for key, value in data.items():
                try:
                    json.dumps(value)
                except TypeError:
                    logger.error(f"Problematic key: {key}, value type: {type(value)}")
                    data[key] = str(value)
        elif isinstance(data, list):
            for i, item in enumerate(data):
                try:
                    json.dumps(item)
                except TypeError:
                    logger.error(f"Problematic item at index {i}, type: {type(item)}")
                    data[i] = str(item)
        # Try again with converted data
        return json.dumps(data)

def init_db():
    """Initialize the database with required tables"""
    try:
        with db_pool.get_connection() as conn:
            cursor = conn.cursor()
            
            # Create tables if they don't exist
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS timetables (
                id TEXT PRIMARY KEY,
                department TEXT NOT NULL,
                programs TEXT NOT NULL,
                levels TEXT NOT NULL,
                session_title TEXT NOT NULL,
                timetable_data TEXT NOT NULL,
                original_data TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            ''')
            
            # Create lecturer_registrations table
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS lecturer_registrations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE,
                department TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                status TEXT DEFAULT 'pending',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                approved_at TIMESTAMP,
                approved_by TEXT
            )
            ''')
            
            # Create users table if it doesn't exist
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL,
                role TEXT NOT NULL,
                department TEXT
            )
            ''')
            
            # Check if department column exists in users table
            cursor.execute("PRAGMA table_info(users)")
            columns = [column[1] for column in cursor.fetchall()]
            
            # Add department column if it doesn't exist
            if 'department' not in columns:
                cursor.execute('ALTER TABLE users ADD COLUMN department TEXT')
            
            # Create indexes for better query performance
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_timetables_department ON timetables(department)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_timetables_session ON timetables(session_title)')
            
            # Add WAL mode for better concurrency if not already set
            cursor.execute('PRAGMA journal_mode=WAL')
            cursor.execute('PRAGMA synchronous=NORMAL')
            cursor.execute('PRAGMA cache_size=-2000')  # 2MB cache
            
            conn.commit()
            
            # Create approvals table if it doesn't exist
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS approvals (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    registration_id TEXT NOT NULL,
                    approved_at TIMESTAMP,
                    approved_by TEXT,
                    FOREIGN KEY (registration_id) REFERENCES timetables(registration_id)
                )
            ''')
            
            conn.commit()
        return True
    except Exception as e:
        logger.error(f"Error initializing database: {str(e)}", exc_info=True)
        return False

def register_lecturer(name, email, department, password_hash):
    """Register a new lecturer"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        # Check if email already exists
        c.execute('SELECT id FROM lecturer_registrations WHERE email = ?', (email,))
        if c.fetchone():
            return False, "Email already registered"
        
        # Insert new registration
        c.execute('''
            INSERT INTO lecturer_registrations (name, email, department, password_hash)
            VALUES (?, ?, ?, ?)
        ''', (name, email, department, password_hash))
        
        conn.commit()
        conn.close()
        return True, "Registration successful. Please wait for admin approval."
    except Exception as e:
        logger.error(f"Error registering lecturer: {str(e)}")
        return False, "Registration failed. Please try again."

def get_pending_registrations():
    """Get all pending lecturer registrations"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        c.execute('''
            SELECT id, name, email, department, created_at
            FROM lecturer_registrations
            WHERE status = 'pending'
            ORDER BY created_at DESC
        ''')
        
        registrations = []
        for row in c.fetchall():
            registrations.append({
                'id': row[0],
                'name': row[1],
                'email': row[2],
                'department': row[3],
                'created_at': row[4]
            })
        
        conn.close()
        return registrations
    except Exception as e:
        logger.error(f"Error getting pending registrations: {str(e)}")
        return []

def approve_lecturer(registration_id, approved_by):
    """Approve a lecturer registration"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        # Get registration details
        c.execute('''
            SELECT name, email, department, password_hash
            FROM lecturer_registrations
            WHERE id = ? AND status = 'pending'
        ''', (registration_id,))
        
        registration = c.fetchone()
        if not registration:
            return False, "Registration not found or already approved"
        
        # Update registration status
        c.execute('''
            UPDATE lecturer_registrations
            SET status = 'approved',
                approved_at = CURRENT_TIMESTAMP,
                approved_by = ?
            WHERE id = ?
        ''', (approved_by, registration_id))
        
        # Add lecturer to users table with the correct password hash
        c.execute('''
            INSERT INTO users (username, password, role, department)
            VALUES (?, ?, 'lecturer', ?)
        ''', (registration[1], registration[3], registration[2]))
        
        conn.commit()
        conn.close()
        return True, "Lecturer approved successfully"
    except Exception as e:
        logger.error(f"Error approving lecturer: {str(e)}")
        return False, "Approval failed. Please try again."

def reject_lecturer(registration_id, approved_by):
    """Reject a lecturer registration"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        c.execute('''
            UPDATE lecturer_registrations
            SET status = 'rejected',
                approved_at = CURRENT_TIMESTAMP,
                approved_by = ?
            WHERE id = ?
        ''', (approved_by, registration_id))
        
        conn.commit()
        conn.close()
        return True, "Lecturer registration rejected"
    except Exception as e:
        logger.error(f"Error rejecting lecturer: {str(e)}")
        return False, "Rejection failed. Please try again."

def validate_timetable_data(timetable_data, department, programs, levels, session_title):
    """Validate timetable data before saving"""
    try:
        if not timetable_data or not isinstance(timetable_data, dict):
            return False, "Invalid timetable data format"
            
        if not timetable_data.get('slots'):
            return False, "Timetable is missing slot data"
            
        # Validate department and session
        if not department or not session_title:
            return False, "Department and session title are required"
            
        # Validate programs and levels
        if not programs or not levels:
            return False, "Programs and levels are required"
            
        # Check for empty timetable
        has_entries = False
        for day in timetable_data['slots']:
            for slot in timetable_data['slots'][day]:
                if timetable_data['slots'][day][slot]:
                    has_entries = True
                    break
            if has_entries:
                                break
        
        if not has_entries:
            return False, "Timetable contains no entries"
            
        return True, "Timetable data is valid"
        
    except Exception as e:
        logger.error(f"Error validating timetable data: {str(e)}")
        return False, f"Validation error: {str(e)}"

def save_timetable(department, programs, levels, session_title, timetable_data, original_data):
    """Save a timetable to the database with improved validation"""
    try:
        # Validate timetable data first
        is_valid, message = validate_timetable_data(timetable_data, department, programs, levels, session_title)
        if not is_valid:
            logger.error(f"Timetable validation failed: {message}")
            return False, message
            
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        # Prepare data for storage
        timetable_json = safe_json_dumps(timetable_data)
        original_json = safe_json_dumps(original_data)
        
        # Save timetable
        c.execute('''
            INSERT INTO timetables (department, programs, levels, session_title, timetable_data, original_data)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (department, json.dumps(programs), json.dumps(levels), session_title, timetable_json, original_json))
        
        timetable_id = c.lastrowid
        
        # Auto-create lecturer accounts
        if 'lecturers' in original_data:
            created = auto_create_lecturer_accounts(original_data['lecturers'], conn)
            logger.info(f"Created {created} lecturer accounts for timetable {timetable_id}")
        
        conn.commit()
        conn.close()

        return True, f"Timetable saved successfully with ID {timetable_id}"
        
    except Exception as e:
        logger.error(f"Error saving timetable: {str(e)}")
        return False, f"Failed to save timetable: {str(e)}"

def get_all_timetables():
    """Retrieve all timetables from the database"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        c.execute('''
            SELECT id, department, programs, levels, session_title, created_at
            FROM timetables
            ORDER BY created_at DESC
        ''')
        
        timetables = []
        for row in c.fetchall():
            timetables.append({
                'id': row[0],
                'department': row[1],
                'programs': json.loads(row[2]),
                'levels': json.loads(row[3]),
                'session_title': row[4],
                'created_at': row[5]
            })
        
        return timetables
    except Exception as e:
        logger.error(f"Error retrieving timetables: {str(e)}")
        return []
    finally:
        conn.close()

def convert_time_slot_format(timetable_data):
    """Convert time slot format in timetable data to match the standardized format"""
    if not timetable_data:
        return timetable_data
        
    # If the data is a string, parse it first
    if isinstance(timetable_data, str):
        try:
            timetable_data = json.loads(timetable_data)
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing timetable data in convert_time_slot_format: {str(e)}")
            return timetable_data
    
    if not isinstance(timetable_data, dict) or 'slots' not in timetable_data:
        return timetable_data
        
    # Define the mapping of old formats to new format
    time_slot_mapping = {
        "0800 – 1000": "08:00-10:00",
        "1000 - 1200": "10:00-12:00",
        "1200 -1400": "12:00-14:00",
        "1400 -1600": "14:00-16:00",
        "08:00-10:00": "08:00-10:00",  
        "10:00-12:00": "10:00-12:00",
        "12:00-14:00": "12:00-14:00",
        "14:00-16:00": "14:00-16:00"
    }
    
    # Convert slots
    converted_slots = {}
    for day, slots in timetable_data['slots'].items():
        converted_slots[day] = {}
        for old_slot, entries in slots.items():
            new_slot = time_slot_mapping.get(old_slot, old_slot)
            converted_slots[day][new_slot] = entries
    
    timetable_data['slots'] = converted_slots
    return timetable_data

def get_timetable_by_id(timetable_id):
    """Get a timetable by its ID"""
    try:
        conn = sqlite3.connect('timetables.db')
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT * FROM timetables WHERE id = ?
        """, (timetable_id,))
        
        row = cursor.fetchone()
        if row:
            timetable_data = row[5]
            original_data = row[6]
            
            if isinstance(timetable_data, str):
                try:
                    timetable_data = json.loads(timetable_data)
                except json.JSONDecodeError as e:
                    logger.error(f"JSON decode error: {str(e)}")
                    return None
            
            # Ensure original_data is a dict
            if isinstance(original_data, str):
                try:
                    original_data = json.loads(original_data)
                except json.JSONDecodeError as e:
                    logger.error(f"JSON decode error (original_data): {str(e)}")
                    return None
                
            # Convert timetable data while keeping original
            converted_data = convert_time_slot_format(timetable_data)
            
            timetable = {
                'id': row[0],
                'department': row[1],
                'programs': json.loads(row[2]) if isinstance(row[2], str) else row[2],
                'levels': json.loads(row[3]) if isinstance(row[3], str) else row[3],
                'session_title': row[4],
                'timetable_data': converted_data,
                'original_data': original_data,
                'created_at': row[7],
                'docx_data': row[8],
                'original_timetable_data': timetable_data  # Keep original data
            }
            return timetable
        return None
    except Exception as e:
        logger.error(f"Error getting timetable: {str(e)}")
        return None
    finally:
        conn.close()

def delete_timetable(timetable_id):
    """Delete a timetable from the database"""
    conn = sqlite3.connect('timetables.db')
    c = conn.cursor()
    
    c.execute('DELETE FROM timetables WHERE id = ?', (timetable_id,))
    
    conn.commit()
    conn.close() 

def get_student_timetable(student_id, program, level):
    """Get a student's timetable with improved error handling and validation"""
    try:
        logger.info(f"Starting timetable generation for {program} Level {level}")
        
        # Get the most recent timetable
        timetables = get_all_timetables()
        if not timetables:
            logger.error("No timetables found in database")
            return None, None, "No timetables found in database"
            
        latest_timetable = get_timetable_by_id(timetables[0]['id'])
        if not latest_timetable:
            logger.error("Failed to retrieve latest timetable")
            return None, None, "Failed to retrieve latest timetable"
            
        logger.debug(f"Retrieved timetable: {latest_timetable}")
        
        # Get timetable data and original data
        timetable_data = latest_timetable.get('timetable_data')  # Use timetable_data for slots
        original_data = latest_timetable.get('original_data')  # Use original_data for module definitions
        if not timetable_data or not original_data:
            logger.error("Required timetable data not found")
            return None, None, "Error: Required timetable data not found"
            
        if isinstance(timetable_data, str):
            try:
                timetable_data = json.loads(timetable_data)
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing timetable data: {str(e)}")
                return None, None, "Error parsing timetable data"
                
        if isinstance(original_data, str):
            try:
                original_data = json.loads(original_data)
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing original data: {str(e)}")
                return None, None, "Error parsing original data"
        
        logger.debug(f"Timetable data structure: {timetable_data}")
        logger.debug(f"Original data structure: {original_data}")
        
        # Validate that this program and level are valid for the timetable
        if program not in latest_timetable['programs']:
            return None, None, f"Program '{program}' is not valid for this timetable"
        if level not in latest_timetable['levels']:
            return None, None, f"Level '{level}' is not valid for this timetable"
            
        # Get all modules targeted for this program and level from original_data
        modules = original_data.get('modules', [])
        relevant_modules = []
        for module in modules:
            target_groups = module.get('target_groups', [])
            # Accept both tuple and list forms
            if any(tuple(group) == (program, level) for group in target_groups):
                relevant_modules.append(module['code'])
        
        if not relevant_modules:
            logger.error(f"No modules found for {program} Level {level}")
            return None, None, f"No modules found for {program} Level {level}"
            
        # Create a filtered timetable for the student
        filtered_timetable = {
            'slots': {day: {slot: [] for slot in TIME_SLOTS} for day in DAYS}
        }
        
        # Process timetable entries
        has_entries = False
        # Check slots data
        logger.debug("Processing slots data...")
        for day in DAYS:
            for slot in TIME_SLOTS:
                entries = timetable_data['slots'].get(day, {}).get(slot, [])
                logger.debug(f"Checking {day} {slot}: {entries}")
                # Handle list of entries
                if isinstance(entries, list):
                    for entry in entries:
                        # Check if this entry is for the student's program and level
                        if entry.get('module') in relevant_modules:
                            logger.debug(f"Found valid entry for {entry.get('module')}: {entry}")
                            # Ensure entry has all required fields
                            if 'lecturer' not in entry:
                                entry['lecturer'] = 'TBA'
                            if 'room' not in entry:
                                entry['room'] = 'TBA'
                            # Prevent duplicates
                            if not any(e['module'] == entry['module'] for e in filtered_timetable['slots'][day][slot]):
                                filtered_timetable['slots'][day][slot].append(entry)
                                has_entries = True
                # Handle single entry
                elif isinstance(entries, dict) and entries.get('module') in relevant_modules:
                    logger.debug(f"Found valid single entry for {entries.get('module')}: {entries}")
                    # Ensure entry has all required fields
                    if 'lecturer' not in entries:
                        entries['lecturer'] = 'TBA'
                    if 'room' not in entries:
                        entries['room'] = 'TBA'
                    # Prevent duplicates
                    if not any(e['module'] == entries['module'] for e in filtered_timetable['slots'][day][slot]):
                        filtered_timetable['slots'][day][slot].append(entries)
                        has_entries = True
        # Check lecturer slots
        logger.debug("Processing lecturer slots...")
        lecturer_slots = timetable_data.get('lecturer_slots', {})
        for lecturer_id, slots in lecturer_slots.items():
            for day in DAYS:
                for slot in TIME_SLOTS:
                    lecturer_entry = slots.get(day, {}).get(slot)
                    logger.debug(f"Checking lecturer {day} {slot}: {lecturer_entry}")
                    if lecturer_entry in relevant_modules:
                        entry = {
                            'module': lecturer_entry,
                            'lecturer': 'TBA',
                            'room': 'TBA'
                        }
                        # Prevent duplicates
                        if not any(e['module'] == entry['module'] for e in filtered_timetable['slots'][day][slot]):
                            filtered_timetable['slots'][day][slot].append(entry)
                            has_entries = True
        # Check room slots
        logger.debug("Processing room slots...")
        room_slots = timetable_data.get('room_slots', {})
        for room, slots in room_slots.items():
            for day in DAYS:
                for slot in TIME_SLOTS:
                    room_entry = slots.get(day, {}).get(slot)
                    logger.debug(f"Checking room {day} {slot}: {room_entry}")
                    if room_entry in relevant_modules:
                        entry = {
                            'module': room_entry,
                            'lecturer': 'TBA',
                            'room': room
                        }
                        # Prevent duplicates
                        if not any(e['module'] == entry['module'] for e in filtered_timetable['slots'][day][slot]):
                            filtered_timetable['slots'][day][slot].append(entry)
                            has_entries = True
        # Debug logging of all entries found
        logger.debug("Entries found in filtered timetable:")
        for day in DAYS:
            for slot in TIME_SLOTS:
                entries = filtered_timetable['slots'][day][slot]
                if entries:
                    logger.debug(f"Entries found for {day} {slot}: {entries}")
        if not has_entries:
            # Log the raw data for debugging
            logger.warning(f"No classes found for {program} Level {level}")
            logger.debug("Lecturer slots raw data:")
            for lecturer_id, slots in timetable_data.get('lecturer_slots', {}).items():
                logger.debug(f"Lecturer {lecturer_id} slots: {slots}")
            logger.debug(f"Room slots raw data: {timetable_data.get('room_slots', {})}")
            logger.debug(f"Filtered timetable: {filtered_timetable}")
            logger.debug(f"Modules: {timetable_data.get('modules', [])}")
            return None, None, f"No classes found for {program} Level {level}. Please verify your program and level."
        logger.debug(f"Final filtered timetable: {filtered_timetable}")
        # Create DOCX document in combined timetable format
        doc = Document()
        doc.add_heading('Student Timetable', level=1)
        doc.add_paragraph(f'Student ID: {student_id}')
        doc.add_paragraph(f'Program: {program}')
        doc.add_paragraph(f'Level: {level}')
        doc.add_paragraph(f'Department: {latest_timetable["department"]}')
        doc.add_paragraph(f'Session: {latest_timetable["session_title"]}')
        doc.add_heading('Weekly Schedule', level=2)
        table = doc.add_table(rows=len(DAYS)+1, cols=len(TIME_SLOTS)+1)
        table.style = 'Table Grid'
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Day"
        for col_idx, slot in enumerate(TIME_SLOTS, 1):
            header_cells[col_idx].text = slot
        # Add data
        for row_idx, day in enumerate(DAYS, 1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = day
            for col_idx, slot in enumerate(TIME_SLOTS, 1):
                try:
                    entries = filtered_timetable['slots'][day][slot]
                    if entries:
                        cell_text = []
                        for entry in entries:
                            cell_text.append(f"{entry['module']}\nRoom: {entry['room']}\nLecturer: {entry['lecturer']}")
                        row_cells[col_idx].text = '\n\n'.join(cell_text)
                    else:
                        row_cells[col_idx].text = ''
                except KeyError as e:
                    row_cells[col_idx].text = 'Error'
                except Exception as e:
                    row_cells[col_idx].text = 'Error'
        # Add important notes section
        doc.add_heading('Important Notes', level=2)
        notes = [
            "1. Please arrive at least 5 minutes before each class.",
            "2. Bring your student ID card to all classes.",
            "3. Notify your lecturer in advance if you need to miss a class.",
            "4. Check the university website regularly for timetable changes.",
            "5. Contact your department office for any timetable-related queries."
        ]
        for note in notes:
            doc.add_paragraph(note, style='List Bullet')
        # Save to bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return filtered_timetable, docx_buffer, None
    except Exception as e:
        logger.error(f"Error in get_student_timetable: {str(e)}", exc_info=True)
        return None, None, f"Error generating timetable: {str(e)}"

def get_departments():
    """Get list of available departments"""
    return [
        "Computer Science",
        "Computer Systems Engineering",
        "Software Engineering",
        "Computer Security",
        "Information Systems",
        "Information Technology",
        "Mathematics",
        "Statistics",
        "Physics",
        "Chemistry",
        "Biology",
        "Geography",
        "Environmental Science",
        "Agriculture",
        "Economics",
        "Accounting",
        "Business Management",
        "Marketing",
        "Human Resources",
        "Law",
        "Psychology",
        "Sociology",
        "History",
        "English",
        "African Languages",
        "Media Studies",
        "Journalism",
        "Public Relations",
        "Political Science",
        "International Relations"
    ] 

def send_lecturer_account_email(email, username, password):
    """Send email notification to lecturer about their account"""
    try:
        # SMTP configuration for sending lecturer account emails
        smtp_server = 'smtp.gmail.com'
        smtp_port = 587
        sender_email = 'matthewkasango@gmail.com'
        sender_password = 'wfxg alzt djek mjgp'
        
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = email
        msg['Subject'] = 'Your MSU Timetable System Account'
        
        body = f"""
        Dear {username},
        
        Your lecturer account has been created in the MSU Timetable System.
        
        Login Details:
        Username: {username}
        Password: {password}
        
        Please change your password after your first login.
        
        Best regards,
        MSU Timetable System
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, email, msg.as_string())
        server.quit()
        
        return True
    except Exception as e:
        logger.error(f"Error sending lecturer account email: {str(e)}")
        return False

def auto_create_lecturer_accounts(lecturers, conn=None):
    """Create lecturer accounts with improved error handling and email notifications"""
    try:
        if not conn:
            conn = sqlite3.connect('timetables.db')
        c = conn.cursor()

        created_count = 0
        for lecturer in lecturers:
            try:
                # Check if lecturer already exists
                c.execute('SELECT username FROM users WHERE username = ?', (lecturer['name'],))
                if c.fetchone():
                    logger.info(f"Lecturer account already exists: {lecturer['name']}")
                    continue

                # Generate default password
                default_password = 'mathy11'
                password_hash = generate_password_hash(default_password)

                # Create lecturer account
                c.execute('''
                    INSERT INTO users (username, password, role, department)
                    VALUES (?, ?, 'lecturer', ?)
                ''', (lecturer['name'], password_hash, lecturer.get('department', 'Unknown')))

                # Send email notification if email is available
                if lecturer.get('email'):
                    if send_lecturer_account_email(lecturer['email'], lecturer['name'], default_password):
                        logger.info(f"Sent account email to {lecturer['name']}")
                    else:
                        logger.warning(f"Failed to send account email to {lecturer['name']}")

                created_count += 1

            except Exception as e:
                logger.error(f"Error creating account for lecturer {lecturer.get('name', 'Unknown')}: {str(e)}")
                continue

        if not conn:
            conn.commit()
            conn.close()
        else:
            conn.commit()

        return created_count

    except Exception as e:
        logger.error(f"Error in auto_create_lecturer_accounts: {str(e)}")
        if conn:
            conn.close()
        return 0

def sync_lecturer_accounts_from_timetables(default_password="mathy11"):
    """Sync lecturer accounts from timetable data"""
    # First ensure database schema is up to date
    init_db()
    
    # Connect to both databases
    timetables_conn = sqlite3.connect('timetables.db')
    users_conn = sqlite3.connect('users.db')
    timetables_c = timetables_conn.cursor()
    users_c = users_conn.cursor()
    
    # Use the same hashing method as login.py
    def hash_password(password):
        return hashlib.sha256(password.encode()).hexdigest()
    
    default_password_hash = hash_password(default_password)
    
    # Debug: Count total timetables
    timetables_c.execute("SELECT COUNT(*) FROM timetables")
    total_timetables = timetables_c.fetchone()[0]
    print(f"Total timetables found: {total_timetables}")
    
    timetables_c.execute("SELECT original_data FROM timetables")
    all_timetables = timetables_c.fetchall()
    created = 0
    total_lecturers_processed = 0
    
    for (original_data_json,) in all_timetables:
        try:
            original_data = json.loads(original_data_json)
            lecturers = original_data.get('lecturers', [])
            print(f"Found {len(lecturers)} lecturers in timetable")
            total_lecturers_processed += len(lecturers)
            
            for lecturer in lecturers:
                username = lecturer.get('email', lecturer['name'])
                department = lecturer.get('department', '')
                print(f"Processing lecturer: {username} from {department}")
                
                # Check if user already exists in users.db
                users_c.execute("SELECT 1 FROM users WHERE username = ?", (username,))
                if not users_c.fetchone():
                    print(f"Creating new account for {username}")
                    users_c.execute(
                        "INSERT INTO users (username, password, role, email) VALUES (?, ?, 'lecturer', ?)",
                        (username, default_password_hash, username)  # Using username as email if no email provided
                    )
                    created += 1
                else:
                    print(f"Account already exists for {username}")
        except Exception as e:
            print(f"Error processing timetable: {str(e)}")
            continue
    
    users_conn.commit()
    users_conn.close()
    timetables_conn.close()
    print(f"Total lecturers processed: {total_lecturers_processed}")
    print(f"New accounts created: {created}")
    return created

def get_all_lecturer_accounts():
    """Get all lecturer accounts from the users table"""
    conn = sqlite3.connect('timetables.db')
    c = conn.cursor()
    
    # First check what columns exist in the users table
    c.execute("PRAGMA table_info(users)")
    columns = [column[1] for column in c.fetchall()]
    print(f"Available columns in users table: {columns}")
    
    # Build the SELECT statement based on available columns
    select_columns = []
    if 'username' in columns:
        select_columns.append('username')
    if 'department' in columns:
        select_columns.append('department')
    if 'role' in columns:
        select_columns.append('role')
    
    if not select_columns:
        print("No valid columns found in users table")
        return []
    
    query = f'''
        SELECT {', '.join(select_columns)}
        FROM users 
        WHERE role = 'lecturer'
        ORDER BY department, username
    '''
    
    c.execute(query)
    lecturers = []
    for row in c.fetchall():
        lecturer = {}
        for i, col in enumerate(select_columns):
            lecturer[col] = row[i]
        lecturers.append(lecturer)
    
    conn.close()
    return lecturers

def get_lecturer_statistics():
    """
    Aggregate scheduled and required hours for each lecturer, including student groups for each slot.
    Returns a list of dicts: [{lecturer, scheduled_hours, required_hours, status, slots: [...]}, ...]
    """
    import sqlite3
    import logging
    from constants import DAYS, TIME_SLOTS, FACULTY_MIN_HOURS, FACULTY_MAX_HOURS
    logger = logging.getLogger(__name__)
    conn = None
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        c.execute('SELECT timetable_data, department, session_title, lecturer_name FROM lecturer_timetables ORDER BY created_at DESC')
        lecturer_stats = {}
        # To get student groups, we need module definitions from the latest timetable
        c2 = conn.cursor()
        c2.execute('SELECT original_data FROM timetables ORDER BY created_at DESC LIMIT 1')
        modules = []
        module_map = {}
        row2 = c2.fetchone()
        if row2:
            try:
                original_data = json.loads(row2[0])
                modules = original_data.get('modules', [])
                module_map = {m['code']: m for m in modules if 'code' in m}
            except Exception as e:
                logger.error(f"Failed to parse original_data for modules: {e}")
        for row in c.fetchall():
            timetable_data, department, session_title, lecturer_name = row
            try:
                timetable = json.loads(timetable_data)
            except Exception as e:
                logger.error(f"Failed to parse timetable for {lecturer_name}: {e}")
                continue
            scheduled_hours = 0
            slots_list = []
            for day in DAYS:
                for slot in TIME_SLOTS:
                    entries = timetable['slots'].get(day, {}).get(slot, [])
                    if isinstance(entries, dict):
                        entries = [entries] if entries else []
                    for entry in entries:
                        scheduled_hours += 2  # each slot = 2 hours
                        module_code = entry.get('module')
                        groups = []
                        module = module_map.get(module_code)
                        if module and 'target_groups' in module:
                            for target_group in module['target_groups']:
                                if isinstance(target_group, (list, tuple)) and len(target_group) == 2:
                                    prog, lvl = target_group
                                    groups.append(f"{prog} {lvl}")
                        slots_list.append({
                            'day': day,
                            'slot': slot,
                            'module': module_code,
                            'room': entry.get('room'),
                            'groups': groups
                        })
            # Use constants for required hours
            required_hours = FACULTY_MIN_HOURS
            status = 'complete' if scheduled_hours >= required_hours else 'incomplete'
            lecturer_stats[lecturer_name] = {
                'lecturer': lecturer_name,
                'department': department,
                'session_title': session_title,
                'scheduled_hours': scheduled_hours,
                'required_hours': required_hours,
                'status': status,
                'slots': slots_list
            }
            logger.info(f"Lecturer: {lecturer_name}, Scheduled: {scheduled_hours}, Required: {required_hours}, Status: {status}")
        return list(lecturer_stats.values())
    except Exception as e:
        logger.error(f"Error in get_lecturer_statistics: {e}")
        return []
    finally:
        if conn:
            conn.close()


def save_lecturer_timetable(lecturer_name, timetable_data, department, session_title):
    """Save a lecturer's timetable to the database"""
    try:
        conn = sqlite3.connect('timetables.db')
        c = conn.cursor()
        
        # Create lecturer_timetables table if it doesn't exist
        c.execute('''
            CREATE TABLE IF NOT EXISTS lecturer_timetables (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                lecturer_name TEXT NOT NULL,
                timetable_data TEXT NOT NULL,
                department TEXT NOT NULL,
                session_title TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Save the timetable
        c.execute('''
            INSERT INTO lecturer_timetables (lecturer_name, timetable_data, department, session_title)
            VALUES (?, ?, ?, ?)
        ''', (lecturer_name, json.dumps(timetable_data), department, session_title))
        
        conn.commit()
        logger.info(f"Saved timetable for lecturer {lecturer_name}")
        return True
    except Exception as e:
        logger.error(f"Error saving lecturer timetable: {str(e)}")
        return False
    finally:
        if conn:
            conn.close()

def get_lecturer_timetable(lecturer_name, timetable_id):
    """Return a filtered timetable (and a DOCX buffer) for the given lecturer (from the timetable record with id timetable_id)."""
    try:
        # Fetch the timetable record (using get_timetable_by_id) – do not rely on extra columns (e.g. created_at)
        timetable_record = get_timetable_by_id(timetable_id)
        if not timetable_record:
            return (None, None, "Timetable record not found.")
        timetable_data = timetable_record["timetable_data"]
        original_data = timetable_record["original_data"]
        logger.debug(f"Fetched timetable record (id {timetable_id}) for lecturer {lecturer_name}.")
        # (Optional: log extra info, e.g. session_title, department, etc.)
        session_title = timetable_record.get("session_title", "Unknown Session")
        department = timetable_record.get("department", "Unknown Department")
        logger.debug(f"Timetable session: {session_title}, department: {department}.")
        # (End optional logging.)
        # Build a filtered timetable (using a deep copy of timetable_data) so that only slots (for each day and time slot) that contain an entry with lecturer_name are kept.
        import copy
        filtered_timetable = copy.deepcopy(timetable_data)
        has_entries = False
        for day in DAYS:
            for slot in TIME_SLOTS:
                # (Assume that timetable_data["slots"][day][slot] is a list (or a dict) of entries.)
                entries = filtered_timetable["slots"][day][slot]
                if isinstance(entries, list):
                    filtered_entries = [e for e in entries if e and e.get("lecturer") == lecturer_name]
                    if filtered_entries:
                        has_entries = True
                    filtered_timetable["slots"][day][slot] = filtered_entries
                elif isinstance(entries, dict) and entries and entries.get("lecturer") == lecturer_name:
                    has_entries = True
                else:
                    filtered_timetable["slots"][day][slot] = []
        if not has_entries:
            logger.warning("No classes found for lecturer {lecturer_name} in timetable (id {timetable_id}).")
            return (None, None, "No classes found for lecturer {lecturer_name} in timetable (id {timetable_id}).")
        # (Optional: log a summary of the filtered timetable.)
        logger.debug("Filtered timetable (for lecturer {lecturer_name}) built successfully.")
        # (End optional logging.)
        # (Optional: build a DOCX (using docx.Document) for download, e.g. as in get_student_timetable.)
        from docx import Document
        import io
        doc = Document()
        doc.add_heading("Lecturer Timetable", level=1)
        doc.add_paragraph(f"Lecturer: {lecturer_name} (Timetable ID: {timetable_id})")
        doc.add_paragraph(f"Session: {session_title} (Department: {department})")
        doc.add_heading("Weekly Schedule", level=2)
        table = doc.add_table(rows=len(DAYS) + 1, cols=len(TIME_SLOTS) + 1)
        table.style = "Table Grid"
        # (Header row: "Day" and time slots.)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Day"
        for (i, slot) in enumerate(TIME_SLOTS, 1):
            header_cells[i].text = slot
        # (Body rows: for each day, fill cells (using filtered_timetable) – if a cell is empty, leave it blank.)
        for (i, day) in enumerate(DAYS, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = day
            for (j, slot) in enumerate(TIME_SLOTS, 1):
                cell_entries = filtered_timetable["slots"][day][slot]
                if cell_entries:
                    cell_text = []
                    for entry in cell_entries:
                        cell_text.append(f"{entry['module']} (Room: {entry['room']})")
                    row_cells[j].text = "\n".join(cell_text)
                else:
                    row_cells[j].text = ""
        # (Optional: add a "Notes" section (or extra info) at the end of the DOCX.)
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(f"This timetable is generated for lecturer {lecturer_name} (Timetable ID: {timetable_id}).")
        lecturer_notes = [
            "Please arrive at least 5 minutes before each class.",
            "Bring your staff ID card to all classes.",
            "Notify your department head in advance if you need to miss a class.",
            "Check the university website regularly for any timetable changes.",
            "Contact your department office for any timetable-related queries.",
            "This timetable is valid for the current academic session only.",
            "All classes are 2 hours long.",
            "Report any timetable conflicts to your department office immediately.",
            "Keep this timetable with you at all times."
        ]
        for note in lecturer_notes:
            doc.add_paragraph(note, style='List Bullet')
        # (Save the DOCX into a BytesIO buffer.)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        # (Return (filtered_timetable, docx_buffer, None) on success.)
        return (filtered_timetable, docx_buffer, None)
    except Exception as e:
        logger.error(f"Error generating lecturer timetable: {e}", exc_info=True)
        return (None, None, f"Error generating lecturer timetable: {e}")

def show_audit_log_table():
    # Implementation of show_audit_log_table function
    pass